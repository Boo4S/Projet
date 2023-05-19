import requests
import openpyxl
import sys
import time
import re
import spacy
nlp = spacy.load("fr_core_news_sm")
from PySide6.QtCore import Qt
from PySide6.QtGui import QPalette, QColor, QStandardItemModel, QStandardItem
from PySide6.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget, QStackedWidget, QFrame, QPushButton, QLabel, QSizePolicy, QComboBox, QCheckBox
from qtawesome import icon
from PySide6.QtCore import Qt, QDate
from PySide6.QtWidgets import QFileDialog , QTreeView, QGroupBox, QRadioButton, QListWidget, QProgressBar, QLineEdit, QSpacerItem, QTableView, QTextEdit, QTableWidget, QTableWidgetItem, QDateEdit,QToolBar, QHeaderView
import pandas as pd
import numpy as np
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlibwidget import MatplotlibWidget
from docx import Document
from openpyxl import Workbook
import os
from datetime import datetime
from dateparser import parse
from langdetect import detect
import os.path
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import nltk
from sentence_transformers import SentenceTransformer
from concurrent_watch import ConcurrentWatch

model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
nltk.download('punkt')
nltk.download('stopwords')

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()


        self.setWindowTitle("Interface RESOSearch")
        self.setGeometry(100, 100, 500, 400)

        self.competitive_intel_widget = QWidget()
        self.setCentralWidget(self.competitive_intel_widget)

        # Création des widgets
        self.central_widget = QWidget()
        self.menu_widget = QFrame()
        self.menu_layout = QVBoxLayout()
        self.stacked_widget = QStackedWidget()
        self.conc_list = QComboBox()
        self.competitive_intel_layout = QVBoxLayout()

        # Configuration du menu
        self.menu_widget.setFixedWidth(200)
        self.menu_widget.setLayout(self.menu_layout)
        self.menu_widget.setFrameShape(QFrame.StyledPanel)

        # Boutons du menu
        button_style = """
        QPushButton {
            color: white;
            background-color: #5F9EA0;
            font-weight: bold;
            border: none;
            padding: 10px;
        }
        QPushButton:hover {
            background-color: #7FB1B5;
        }
        """
        # Boutons du menu
        self.cleaner_button = QPushButton(icon('fa5s.search', color='white'), "RESOSearch")
        self.cleaner_button.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(0))
        self.cleaner_button.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Maximum)
        self.cleaner_button.setStyleSheet(button_style)
        self.menu_layout.addWidget(self.cleaner_button)


        # Bouton Veille concurrentielle
        self.competitive_intel_button = QPushButton(icon('fa5s.industry', color='white'), "Veille concurrentielle")
        self.competitive_intel_button.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(1))
        self.competitive_intel_button.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Maximum)
        self.competitive_intel_button.setStyleSheet(button_style)
        self.menu_layout.addWidget(self.competitive_intel_button)


        # Configuration de la page RESOSearch
        self.cleaner_layout = QVBoxLayout()

        # Création d'un QVBoxLayout pour la première colonne
        self.column1_layout = QVBoxLayout()

        # Ajout de la zone de texte
        self.search_text = QLineEdit()
        self.search_text.setPlaceholderText("Entrez des mots-clés pour rechercher des articles")
        self.search_text.setFixedWidth(200)
        self.search_text.setFixedHeight(30)
        self.search_text.setStyleSheet("color: white; background-color: #353535;")

        self.column1_layout.addWidget(self.search_text, alignment=Qt.AlignTop)

        # Création d'un QVBoxLayout pour les options d'extraction
        self.export_options_layout = QVBoxLayout()

        # Ajout des QCheckBox pour les formats d'extraction
        self.export_excel_checkbox = QCheckBox("Exporter en Excel")
        self.export_options_layout.addWidget(self.export_excel_checkbox)

        self.export_word_checkbox = QCheckBox("Exporter en Word")
        self.export_options_layout.addWidget(self.export_word_checkbox)

        # Création d'un QVBoxLayout pour les sélecteurs de dates
        self.date_range_layout = QVBoxLayout()

        # Ajout d'un QLabel pour la date de début
        self.start_date_label = QLabel("Date de début :")
        self.date_range_layout.addWidget(self.start_date_label)

        # Ajout du sélecteur de date de début
        self.start_date_edit = QDateEdit()
        self.start_date_edit.setDate(QDate.currentDate())
        self.start_date_edit.setCalendarPopup(True)
        self.date_range_layout.addWidget(self.start_date_edit)

        # Ajout d'un QLabel pour la date de fin
        self.end_date_label = QLabel("Date de fin :")
        self.date_range_layout.addWidget(self.end_date_label)

        # Ajout du sélecteur de date de fin
        self.end_date_edit = QDateEdit()
        self.end_date_edit.setDate(QDate.currentDate())
        self.end_date_edit.setCalendarPopup(True)
        self.date_range_layout.addWidget(self.end_date_edit)

        # Ajout du QVBoxLayout des sélecteurs de dates à la première colonne
        self.column1_layout.addLayout(self.date_range_layout)

        # Ajout d'un QSpacerItem pour combler l'espace entre les options d'exportation et le bas de la fenêtre
        spacer = QSpacerItem(0, 0, QSizePolicy.Minimum, QSizePolicy.Expanding)
        self.export_options_layout.addItem(spacer)

        # Ajout du QVBoxLayout des options d'extraction à la première colonne
        self.column1_layout.addLayout(self.export_options_layout)


        # Création d'un QVBoxLayout pour la deuxième colonne
        self.column2_layout = QVBoxLayout()

        # Ajout de la QComboBox pour choisir le navigateur
        self.browser_combobox = QComboBox()
        self.browser_combobox.addItem("Google Chrome")
        self.browser_combobox.addItem("Mozilla Firefox")
        self.browser_combobox.addItem("Microsoft Edge")
        self.browser_combobox.addItem("Safari")
        self.browser_combobox.addItem("Opera")
        self.browser_combobox.setFixedWidth(150)
        self.browser_combobox.setFixedHeight(30)

        self.column2_layout.addWidget(self.browser_combobox)

        # Création d'un QVBoxLayout pour les options de langues
        self.language_options_layout = QVBoxLayout()

        # Ajout d'un QLabel pour expliquer l'option de choix de la langue
        self.language_label = QLabel("Choisissez la langue de la recherche :")
        self.language_options_layout.addWidget(self.language_label)

        # Ajout de la QComboBox pour choisir la langue
        self.country_combobox = QComboBox()
        self.country_combobox.addItem("France")
        self.country_combobox.addItem("Royaume-Uni")
        self.country_combobox.addItem("Espagne")
        self.country_combobox.addItem("Allemagne")
        self.country_combobox.addItem("Italie")
        self.country_combobox.setFixedWidth(150)
        self.country_combobox.setFixedHeight(30)
        self.language_options_layout.addWidget(self.country_combobox)

        self.language_codes = {
        "France": "fr",
        "Royaume-Uni": "en",
        "Espagne": "es",
        "Allemagne": "de",
        "Italie": "it"
        }

        self.google_domains = {
        "France": "google.fr",
        "Royaume-Uni": "google.co.uk",
        "Espagne": "google.es",
        "Allemagne": "google.de",
        "Italie": "google.it"
            # Ajoutez d'autres pays et domaines Google si nécessaire
        }

        self.country_codes = {
        "France": "fr",
        "Royaume-Uni": "uk",
        "Espagne": "es",
        "Allemagne": "de",
        "Italie": "it"
        }

        # Ajout de la QComboBox pour choisir la source de recherche
        self.source_combobox = QComboBox()
        self.source_combobox.addItem("Tous")
        self.source_combobox.addItem("Actualités")
        self.source_combobox.setFixedWidth(150)
        self.source_combobox.setFixedHeight(30)

        # Ajout d'un QLabel pour expliquer l'option de source de recherche
        self.source_label = QLabel("Source de recherche :")
        self.language_options_layout.addWidget(self.source_label)

        # Ajout de la QComboBox des sources de recherche à la deuxième colonne
        self.language_options_layout.addWidget(self.source_combobox)

        # Ajout d'un QSpacerItem pour pousser les options de langues vers le haut
        spacer = QSpacerItem(0, 0, QSizePolicy.Minimum, QSizePolicy.Expanding)
        self.language_options_layout.addItem(spacer)

        # Ajout du QVBoxLayout des options de langues à la deuxième colonne
        self.column2_layout.addLayout(self.language_options_layout)

        # Ajout d'un QHBoxLayout pour contenir les deux colonnes
        self.top_layout = QHBoxLayout()
        self.top_layout.addLayout(self.column1_layout)
        self.top_layout.addLayout(self.column2_layout)

        # Ajout du QHBoxLayout au QVBoxLayout principal
        self.cleaner_layout.insertLayout(0, self.top_layout)

        # Ajout du QHBoxLayout au QVBoxLayout principal
        self.cleaner_layout.insertLayout(0, self.top_layout)

        # Ajout d'un QHBoxLayout pour contenir la zone de texte et le bouton de recherche
        self.top_layout = QHBoxLayout()
        self.top_layout.addWidget(self.search_text, alignment=Qt.AlignLeft)

        # Ajout du QHBoxLayout au QVBoxLayout principal
        self.cleaner_layout.insertLayout(0, self.top_layout)

        # Création d'un QHBoxLayout pour le bouton de recherche et la barre de progression
        self.button_progress_layout = QHBoxLayout()

        # Ajout d'un espace pour centrer le bouton de recherche
        left_spacer = QSpacerItem(0, 0, QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.button_progress_layout.addItem(left_spacer)

       # Ajout du bouton de recherche
        self.search_button = QPushButton("Rechercher")
        self.search_button.setFixedSize(200, 30)
        self.button_progress_layout.addWidget(self.search_button) 
        self.search_button.clicked.connect(self.perform_search)

        # Ajout d'un espace pour centrer le bouton de recherche
        right_spacer = QSpacerItem(0, 0, QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.button_progress_layout.addItem(right_spacer)

        # Ajout du QHBoxLayout du bouton de recherche à la QVBoxLayout principale
        self.cleaner_layout.addLayout(self.button_progress_layout)

        # Ajout de la barre de progression
        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(100)
        self.progress_bar.setValue(0)
        self.cleaner_layout.addWidget(self.progress_bar)

        self.cleaner_widget = QWidget()
        self.cleaner_widget.setLayout(self.cleaner_layout)
        self.stacked_widget.addWidget(self.cleaner_widget)

        # Création de la page Competitive Intelligence
        self.competitive_intel_layout = QVBoxLayout()

        self.conc_label = QLabel("Concurrents :")
        self.competitive_intel_layout.addWidget(self.conc_label)

        self.conc_list = QComboBox()
        self.competitive_intel_layout.addWidget(self.conc_list)

        self.search_label = QLabel("Ajouter un concurrent :")
        self.competitive_intel_layout.addWidget(self.search_label)

        self.search_input_layout = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input_layout.addWidget(self.search_input)

        self.add_button = QPushButton("Ajouter")
        self.search_input_layout.addWidget(self.add_button)

        self.remove_button = QPushButton("Supprimer")
        self.search_input_layout.addWidget(self.remove_button)

        self.competitive_intel_layout.addLayout(self.search_input_layout)

        self.news_table = QTableWidget()
        self.news_table.setRowCount(10)
        self.news_table.setColumnCount(3)
        self.news_table.setHorizontalHeaderLabels(["Titre", "Date", "Source"])
        header = self.news_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.competitive_intel_layout.addWidget(self.news_table)

        self.competitive_intel_widget.setLayout(self.competitive_intel_layout)

        self.concurrent_watch = ConcurrentWatch(self.conc_list, self.search_input, self.add_button, self.remove_button, self.news_table)

        # Connecter le bouton "Ajouter" à la fonction d'ajout de concurrent
        self.add_button.clicked.connect(self.concurrent_watch.add_concurrent)

        # Connecter le bouton "Supprimer" à la fonction de suppression de concurrent
        self.remove_button.clicked.connect(self.concurrent_watch.remove_concurrent)

        # Ajout de la page de veille concurrentielle au QStackedWidget
        self.competitive_intel_widget = QWidget()
        self.competitive_intel_widget.setLayout(self.competitive_intel_layout)
        self.stacked_widget.addWidget(self.competitive_intel_widget)


        # Configuration du layout principal
        self.main_layout = QHBoxLayout()
        self.main_layout.addWidget(self.menu_widget)

        self.vertical_layout = QVBoxLayout()
        self.vertical_layout.addWidget(self.stacked_widget)
        self.main_layout.addLayout(self.vertical_layout)

        # Configuration du widget central
        self.central_widget.setLayout(self.main_layout)
        self.setCentralWidget(self.central_widget)

        # Ajout de la page ChatGPT à QStackedWidget

    
    def perform_search(self, language_codes):
        # Récupérer les options sélectionnées par l'utilisateur
        search_query = self.search_text.text()
        start_date = self.start_date_edit.date().toString("yyyy-MM-dd")
        end_date = self.end_date_edit.date().toString("yyyy-MM-dd")
        start_date_dt = datetime.strptime(start_date, "%Y-%m-%d")
        end_date_dt = datetime.strptime(end_date, "%Y-%m-%d")
        country = self.country_combobox.currentText()

        common_words = {
            "France": ["le", "la", "les", "un", "une", "des"],
            "Royaume-Uni": ["the", "a", "an", "of", "in", "on"],
            "Espagne": ["el", "la", "los", "las", "un", "una"],
            "Allemagne": ["der", "die", "das", "ein", "eine", "eines"],
            "Italie": ["il", "la", "lo", "gli", "un", "una"]
        }

        # Utilisez la langue du pays sélectionné comme langue de recherche
        language_code = self.language_codes[country]
        country_code = self.country_codes[country]

        # Récupérer la source de recherche sélectionnée
        selected_source = self.source_combobox.currentText()

        # Définir le paramètre tbm en fonction de la source sélectionnée
        source_options = {
            "Tous": None,
            "Actualités": "nws"
        }

        tbm = source_options[selected_source]

        start_index = 0
        has_results = True

        # Utiliser l'API SerpApi pour effectuer la recherche
        extracted_data = []
        while has_results and start_index < 50:
            api_key = "9b0d4c0366546a7bd81c14d13ae3f304ea744bff2faa67fab9eed518194b7f40"
            params = {
                "q": search_query,
                'start': start_index,
                "tbm": tbm,
                "api_key": api_key,
                "hl": language_code,
                "gl": country_code,
                "google_domain": self.google_domains[country],
                'per_page': 100,
                "date_restrict": f"{start_date_dt.strftime('%Y-%m-%d')}..{end_date_dt.strftime('%Y-%m-%d')}"
            }

            if tbm is not None:
                params["tbm"] = tbm

            response = requests.get("https://serpapi.com/search", params=params)
            search_results = response.json()
            print(search_results)

            # Traitez les résultats de la recherche
            if selected_source == "Actualités":
                results_key = "news_results"
            else:
                results_key = "organic_results"

            search_results_list = search_results.get(results_key, [])

            for article in search_results_list:
                publication_date = article.get("date", "")

                # Convertir la date de publication en objet datetime
                if publication_date:
                    publication_date_dt = parse(publication_date)
                else:
                    continue

                # Vérifiez si la date de publication est dans la plage de dates définie
                if start_date_dt <= publication_date_dt <= end_date_dt:
                    title = article.get("title", "")
                    link = article.get("link", "")
                    snippet = article.get("snippet", "")

                    # Ajouter un filtre pour vérifier si l'article est dans la langue sélectionnée
                    detected_language = detect(title)
                    if detected_language != self.language_codes[country]:
                        continue

                    extracted_data.append((publication_date, title, link, snippet))

            # Vérifiez s'il y a encore des résultats à récupérer
            if len(search_results_list) == 0:
                break

            start_index += 10  # Incrémentez start_index pour la prochaine page

            # Mettre à jour la barre de progression
            self.progress_bar.setValue(100)

        # Vous pouvez également exporter les résultats au format Excel ou Word, en utilisant les options sélectionnées
        if self.export_excel_checkbox.isChecked():
            timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")
            excel_file_name = f"{search_query}_resultats_exportes_{timestamp}.xlsx"
            self.export_to_excel(extracted_data, excel_file_name)

        if self.export_word_checkbox.isChecked():
            timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")
            word_file_name = f"{search_query}_resultats_exportes_{timestamp}.docx"
            self.export_to_word(extracted_data, word_file_name)

        # Ajoutez cette ligne à la fin de la méthode perform_search
        return extracted_data

    def create_export_directory(self):
        date_str = datetime.now().strftime("%Y-%m-%d")
        directory = os.path.join("exports", date_str)

        if not os.path.exists(directory):
            os.makedirs(directory)

        return directory

    def export_to_excel(self, data, search_query):
        export_directory = self.create_export_directory()
        timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")

        workbook = Workbook()
        sheet = workbook.active

        # Ajouter les en-têtes de colonnes
        headers = ["Date de publication", "Titre", "Lien", "Résumé"]
        for col_num, header in enumerate(headers, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = header

        # Ajouter les données dans les cellules
        for row_num, row_data in enumerate(data, 2):
            for col_num, cell_data in enumerate(row_data, 1):
                cell = sheet.cell(row=row_num, column=col_num)
                cell.value = cell_data

        # Remplacez les espaces et les caractères spéciaux par des underscores
        sanitized_query = re.sub(r"[^\w\s]", "", search_query).replace(" ", "_")

        # Enregistrer le fichier Excel dans le dossier d'exportation
        excel_file_name = f"{sanitized_query}_resultats_exportes_{timestamp}.xlsx"
        workbook.save(os.path.join(export_directory, excel_file_name))

    def export_to_word(self, data, search_query):
        export_directory = self.create_export_directory()
        timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")
        document = Document()

        # Ajouter les données au document
        for row_data in data:
            date_publication, titre, lien, resume = row_data

            # Ajouter la date de publication
            date_paragraph = document.add_paragraph()
            date_paragraph.add_run(f"Date de publication : {date_publication}").bold = True

            # Ajouter le titre
            document.add_heading(titre, level=1)

            # Ajouter le lien
            document.add_paragraph(f"Lien : {lien}")

            # Ajouter le résumé
            document.add_paragraph(f"Résumé : {resume}")

            # Ajouter un saut de ligne pour séparer les articles
            document.add_paragraph()

        # Remplacez les espaces et les caractères spéciaux par des underscores
        sanitized_query = re.sub(r"[^\w\s]", "", search_query).replace(" ", "_")

        # Enregistrer le fichier Word dans le dossier d'exportation
        word_file_name = f"{sanitized_query}_resultats_exportes_{timestamp}.docx"
        document.save(os.path.join(export_directory, word_file_name))


if __name__ == "__main__":
    app = QApplication(sys.argv)

    # Configuration de la palette de couleurs
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(53, 53, 53))
    palette.setColor(QPalette.WindowText, Qt.white)
    palette.setColor(QPalette.Base, QColor(25, 25, 25))
    palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
    palette.setColor(QPalette.ToolTipBase, Qt.white)
    palette.setColor(QPalette.ToolTipText, Qt.white)
    palette.setColor(QPalette.Text, Qt.white)
    palette.setColor(QPalette.Button, QColor(53, 53, 53))
    palette.setColor(QPalette.ButtonText, Qt.white)
    palette.setColor(QPalette.BrightText, Qt.red)
    palette.setColor(QPalette.Highlight, QColor(142, 45, 197).lighter())
    palette.setColor(QPalette.HighlightedText, Qt.black)
    app.setPalette(palette)
    app.setStyle("Fusion")
    

    # Affichage de la fenêtre principale
    mainWin = MainWindow()
    mainWin.show()
    sys.exit(app.exec())