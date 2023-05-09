import requests
import openpyxl
import sys
import time
import re
import spacy
nlp = spacy.load("fr_core_news_sm")
from bs4 import BeautifulSoup
from PySide6.QtCore import Qt
from PySide6.QtGui import QPalette, QColor
from PySide6.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget, QStackedWidget, QFrame, QPushButton, QLabel, QSizePolicy, QComboBox, QCheckBox
from qtawesome import icon
from PySide6.QtCore import QEasingCurve, QPropertyAnimation, QAbstractTableModel, Qt, QDate
from PySide6.QtWidgets import QGraphicsOpacityEffect, QGroupBox, QRadioButton, QListWidget, QProgressBar, QLineEdit, QSpacerItem, QTableView, QTextEdit, QTableWidget, QTableWidgetItem, QDateEdit,QToolBar, QHeaderView
from PySide6.QtGui import QAction, QPainter
from serpapi import GoogleSearch
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlibwidget import MatplotlibWidget
from docx import Document
from openpyxl import Workbook
import os
from datetime import datetime
from collections import Counter
from urllib.parse import urlparse
from countryinfo import CountryInfo

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Interface RESOSearch")
        self.setGeometry(100, 100, 500, 400)
        

        # Création des widgets
        self.central_widget = QWidget()
        self.menu_widget = QFrame()
        self.menu_layout = QVBoxLayout()
        self.stacked_widget = QStackedWidget()

        # Configuration du menu
        self.menu_widget.setFixedWidth(200)
        self.menu_widget.setLayout(self.menu_layout)
        self.menu_widget.setFrameShape(QFrame.StyledPanel)

        # Boutons du menu
        button_style = """
        QPushButton {
            color: white;
            background-color: #5F9EA0;
            font-weight: bold;
            border: none;
            padding: 10px;
        }
        QPushButton:hover {
            background-color: #7FB1B5;
        }
        """
        # Boutons du menu
        self.cleaner_button = QPushButton(icon('fa5s.broom', color='white'), "RESOSearch")
        self.cleaner_button.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(0))
        self.cleaner_button.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Maximum)
        self.cleaner_button.setStyleSheet(button_style)
        self.menu_layout.addWidget(self.cleaner_button)

        self.registry_button = QPushButton(icon('fa5s.archive', color='white'), "Visualisation des données")
        self.registry_button.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(1))
        self.registry_button.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Maximum)
        self.registry_button.setStyleSheet(button_style)
        self.menu_layout.addWidget(self.registry_button)

        # Bouton Veille concurrentielle
        self.competitive_intel_button = QPushButton(icon('fa5s.cogs', color='white'), "Veille concurrentielle")
        self.competitive_intel_button.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(2))
        self.competitive_intel_button.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Maximum)
        self.competitive_intel_button.setStyleSheet(button_style)
        self.menu_layout.addWidget(self.competitive_intel_button)

        # Bouton Veille stratégique
        self.strategic_intel_button = QPushButton(icon('fa5s.cogs', color='white'), "Veille stratégique")
        self.strategic_intel_button.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(3))
        self.strategic_intel_button.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Maximum)
        self.strategic_intel_button.setStyleSheet(button_style)
        self.menu_layout.addWidget(self.strategic_intel_button)


        # Configuration de la page RESOSearch
        self.cleaner_layout = QVBoxLayout()

        # Création d'un QVBoxLayout pour la première colonne
        self.column1_layout = QVBoxLayout()

        # Ajout de la zone de texte
        self.search_text = QLineEdit()
        self.search_text.setPlaceholderText("Entrez des mots-clés pour rechercher des articles")
        self.search_text.setFixedWidth(200)
        self.search_text.setFixedHeight(30)
        self.search_text.setStyleSheet("color: white; background-color: #353535; border: 1px solid white;")

        self.column1_layout.addWidget(self.search_text, alignment=Qt.AlignTop)

        # Création d'un QVBoxLayout pour les options d'extraction
        self.export_options_layout = QVBoxLayout()

        # Ajout des QCheckBox pour les formats d'extraction
        self.export_excel_checkbox = QCheckBox("Exporter en Excel")
        self.export_options_layout.addWidget(self.export_excel_checkbox)

        self.export_word_checkbox = QCheckBox("Exporter en Word")
        self.export_options_layout.addWidget(self.export_word_checkbox)

        self.sort_relevance_checkbox = QCheckBox("Trier par pertinence")
        self.export_options_layout.addWidget(self.sort_relevance_checkbox)

        self.sort_date_checkbox = QCheckBox("Trier par date")
        self.export_options_layout.addWidget(self.sort_date_checkbox)


        # Création d'un QVBoxLayout pour les sélecteurs de dates
        self.date_range_layout = QVBoxLayout()

        # Ajout d'un QLabel pour la date de début
        self.start_date_label = QLabel("Date de début :")
        self.date_range_layout.addWidget(self.start_date_label)

        # Ajout du sélecteur de date de début
        self.start_date_edit = QDateEdit()
        self.start_date_edit.setDate(QDate.currentDate())
        self.start_date_edit.setCalendarPopup(True)
        self.date_range_layout.addWidget(self.start_date_edit)

        # Ajout d'un QLabel pour la date de fin
        self.end_date_label = QLabel("Date de fin :")
        self.date_range_layout.addWidget(self.end_date_label)

        # Ajout du sélecteur de date de fin
        self.end_date_edit = QDateEdit()
        self.end_date_edit.setDate(QDate.currentDate())
        self.end_date_edit.setCalendarPopup(True)
        self.date_range_layout.addWidget(self.end_date_edit)

        # Ajout du QVBoxLayout des sélecteurs de dates à la première colonne
        self.column1_layout.addLayout(self.date_range_layout)

        # Ajout d'un QSpacerItem pour combler l'espace entre les options d'exportation et le bas de la fenêtre
        spacer = QSpacerItem(0, 0, QSizePolicy.Minimum, QSizePolicy.Expanding)
        self.export_options_layout.addItem(spacer)

        # Ajout du QVBoxLayout des options d'extraction à la première colonne
        self.column1_layout.addLayout(self.export_options_layout)


        # Création d'un QVBoxLayout pour la deuxième colonne
        self.column2_layout = QVBoxLayout()

        # Ajout de la QComboBox pour choisir le navigateur
        self.browser_combobox = QComboBox()
        self.browser_combobox.addItem("Google Chrome")
        self.browser_combobox.addItem("Mozilla Firefox")
        self.browser_combobox.addItem("Microsoft Edge")
        self.browser_combobox.addItem("Safari")
        self.browser_combobox.addItem("Opera")
        self.browser_combobox.setFixedWidth(150)
        self.browser_combobox.setFixedHeight(30)

        self.column2_layout.addWidget(self.browser_combobox)

        # Création d'un QVBoxLayout pour les options de langues
        self.language_options_layout = QVBoxLayout()

        # Ajout d'un QLabel pour expliquer l'option de choix de la langue
        self.language_label = QLabel("Choisissez la langue de la recherche :")
        self.language_options_layout.addWidget(self.language_label)

        # Ajout de la QComboBox pour choisir la langue
        self.country_combobox = QComboBox()
        self.country_combobox.addItem("France")
        self.country_combobox.addItem("Royaume-Uni")
        self.country_combobox.addItem("Espagne")
        self.country_combobox.addItem("Allemagne")
        self.country_combobox.addItem("Italie")
        self.country_combobox.setFixedWidth(150)
        self.country_combobox.setFixedHeight(30)
        self.language_options_layout.addWidget(self.country_combobox)

        self.language_codes = {
        "France": "fr",
        "Royaume-Uni": "en",
        "Espagne": "es",
        "Allemagne": "de",
        "Italie": "it"
        }

        self.google_domains = {
        "France": "google.fr",
        "Royaume-Uni": "google.co.uk",
        "Espagne": "google.es",
        "Allemagne": "google.de",
        "Italie": "google.it"
            # Ajoutez d'autres pays et domaines Google si nécessaire
        }

        self.country_codes = {
        "France": "fr",
        "Royaume-Uni": "uk",
        "Espagne": "es",
        "Allemagne": "de",
        "Italie": "it"
        }

        # Ajout de la QComboBox pour choisir la source de recherche
        self.source_combobox = QComboBox()
        self.source_combobox.addItem("Tous")
        self.source_combobox.addItem("Actualités")
        self.source_combobox.setFixedWidth(150)
        self.source_combobox.setFixedHeight(30)

        # Ajout d'un QLabel pour expliquer l'option de source de recherche
        self.source_label = QLabel("Source de recherche :")
        self.language_options_layout.addWidget(self.source_label)

        # Ajout de la QComboBox des sources de recherche à la deuxième colonne
        self.language_options_layout.addWidget(self.source_combobox)

        # Ajout d'un QSpacerItem pour pousser les options de langues vers le haut
        spacer = QSpacerItem(0, 0, QSizePolicy.Minimum, QSizePolicy.Expanding)
        self.language_options_layout.addItem(spacer)

        # Création d'un QVBoxLayout pour les options de partage
        self.share_options_layout = QVBoxLayout()

        # Ajout d'un QLabel pour expliquer l'option de partage
        self.share_label = QLabel("Partager les articles extraits sur :")
        self.share_options_layout.addWidget(self.share_label)

        # Ajout de la QComboBox pour choisir le mode de partage
        self.share_combobox = QComboBox()
        self.share_combobox.addItem("Aucun")
        self.share_combobox.addItem("Facebook")
        self.share_combobox.addItem("Twitter")
        self.share_combobox.addItem("LinkedIn")
        self.share_combobox.addItem("E-mail")
        self.share_combobox.setFixedWidth(150)
        self.share_combobox.setFixedHeight(30)

        self.share_options_layout.addWidget(self.share_combobox)

        # Ajout du QVBoxLayout des options de partage à la deuxième colonne
        self.column2_layout.addLayout(self.share_options_layout)

        # Ajout du QVBoxLayout des options de langues à la deuxième colonne
        self.column2_layout.addLayout(self.language_options_layout)

        # Ajout d'un QHBoxLayout pour contenir les deux colonnes
        self.top_layout = QHBoxLayout()
        self.top_layout.addLayout(self.column1_layout)
        self.top_layout.addLayout(self.column2_layout)

        # Ajout du QHBoxLayout au QVBoxLayout principal
        self.cleaner_layout.insertLayout(0, self.top_layout)

        # Ajout du QHBoxLayout au QVBoxLayout principal
        self.cleaner_layout.insertLayout(0, self.top_layout)

        # Ajout d'un QHBoxLayout pour contenir la zone de texte et le bouton de recherche
        self.top_layout = QHBoxLayout()
        self.top_layout.addWidget(self.search_text, alignment=Qt.AlignLeft)

        # Ajout du QHBoxLayout au QVBoxLayout principal
        self.cleaner_layout.insertLayout(0, self.top_layout)

        # Création d'un QHBoxLayout pour le bouton de recherche et la barre de progression
        self.button_progress_layout = QHBoxLayout()

        # Ajout d'un espace pour centrer le bouton de recherche
        left_spacer = QSpacerItem(0, 0, QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.button_progress_layout.addItem(left_spacer)

       # Ajout du bouton de recherche
        self.search_button = QPushButton("Rechercher")
        self.search_button.setFixedSize(200, 30)
        self.button_progress_layout.addWidget(self.search_button) 
        self.search_button.clicked.connect(self.perform_search)
        self.search_button.clicked.connect(self.on_search_button_clicked)

        # Ajout d'un espace pour centrer le bouton de recherche
        right_spacer = QSpacerItem(0, 0, QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.button_progress_layout.addItem(right_spacer)

        # Ajout du QHBoxLayout du bouton de recherche à la QVBoxLayout principale
        self.cleaner_layout.addLayout(self.button_progress_layout)

        # Ajout de la barre de progression
        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(100)
        self.progress_bar.setValue(0)
        self.cleaner_layout.addWidget(self.progress_bar)

        self.cleaner_widget = QWidget()
        self.cleaner_widget.setLayout(self.cleaner_layout)
        self.stacked_widget.addWidget(self.cleaner_widget)

        # Création de la page Visualisation des données
        self.visualization_layout = QVBoxLayout()

        # Liste déroulante pour sélectionner le type de graphique
        self.chart_type_combobox = QComboBox()
        self.chart_type_combobox.addItems(["Histogramme", "Diagramme à barres", "Diagramme en secteurs"])
        self.visualization_layout.addWidget(self.chart_type_combobox)

        # Options pour sélectionner les données à visualiser
        self.option_group = QGroupBox("Options de visualisation")
        self.option_layout = QVBoxLayout()

        self.option_keyword = QRadioButton("Pertinence des mots-clés")
        self.option_website = QRadioButton("Sites les plus fréquents")
        self.option_country = QRadioButton("Pays avec le plus d'actualités")

        self.option_layout.addWidget(self.option_keyword)
        self.option_layout.addWidget(self.option_website)
        self.option_layout.addWidget(self.option_country)

        self.option_group.setLayout(self.option_layout)
        self.visualization_layout.addWidget(self.option_group)

        # Espace pour afficher le graphique ou le diagramme sélectionné
        self.matplotlib_widget = MatplotlibWidget()
        self.visualization_layout.addWidget(self.matplotlib_widget)

        # Bouton pour exporter le graphique ou le diagramme
        self.export_button = QPushButton("Exporter en image ou PDF")
        self.visualization_layout.addWidget(self.export_button)

        # Ajout de la page de visualisation des données à QStackedWidget
        self.visualization_widget = QWidget()
        self.visualization_widget.setLayout(self.visualization_layout)
        self.stacked_widget.addWidget(self.visualization_widget)

        # Configuration de la page de veille concurrentielle
        self.competitive_intel_layout = QVBoxLayout()

        # Liste des concurrents
        self.conc_label = QLabel("Concurrents :")
        self.competitive_intel_layout.addWidget(self.conc_label)

        self.conc_list = QComboBox()
        self.conc_list.addItem("Concurrent 1")
        self.conc_list.addItem("Concurrent 2")
        self.conc_list.addItem("Concurrent 3")
        self.competitive_intel_layout.addWidget(self.conc_list)

        # Champ de recherche pour ajouter un concurrent
        self.search_label = QLabel("Ajouter un concurrent :")
        self.competitive_intel_layout.addWidget(self.search_label)

        self.search_input = QLineEdit()
        self.competitive_intel_layout.addWidget(self.search_input)

        self.search_button = QPushButton("Ajouter")
        self.competitive_intel_layout.addWidget(self.search_button)

        # Tableau des informations sur les concurrents
        self.info_table = QTableWidget()
        self.info_table.setRowCount(10)
        self.info_table.setColumnCount(3)
        self.info_table.setHorizontalHeaderLabels(["Titre", "Date", "Source"])
        header = self.info_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.competitive_intel_layout.addWidget(self.info_table)

        # Filtres et bouton de rafraîchissement
        self.filters_layout = QHBoxLayout()
        self.keyword_filter = QLineEdit()
        self.keyword_filter.setPlaceholderText("Filtrer par mots-clés")
        self.filters_layout.addWidget(self.keyword_filter)
        self.keyword_filter.setStyleSheet("color: white; background-color: #353535; border:")

        self.source_filter = QLineEdit()
        self.source_filter.setPlaceholderText("Filtrer par source")
        self.filters_layout.addWidget(self.source_filter)
        self.source_filter.setStyleSheet("color: white; background-color: #353535;")
        
        self.refresh_button = QPushButton("Rafraîchir")
        self.filters_layout.addWidget(self.refresh_button)

        self.competitive_intel_layout.addLayout(self.filters_layout)

        # Ajout de la page de veille concurrentielle au QStackedWidget
        self.competitive_intel_widget = QWidget()
        self.competitive_intel_widget.setLayout(self.competitive_intel_layout)
        self.stacked_widget.addWidget(self.competitive_intel_widget)

        # Ajout de la page de veille stratégique au QStackedWidget
        self.strategic_intel_widget = QWidget()
        self.strategic_intel_layout = QVBoxLayout()

        # Création de la liste de tendances, d'évolutions technologiques et d'opportunités de croissance à suivre
        self.trends_list_widget = QListWidget()
        self.trends_list_widget.setFixedSize(200, 180)
        self.trends_list_widget.addItem("Tendance 1")
        self.trends_list_widget.addItem("Tendance 2")
        self.trends_list_widget.addItem("Tendance 3")
        self.trends_list_widget.addItem("Tendance 4")

        # Ajout du champ de recherche pour ajouter de nouvelles tendances ou opportunités à la liste
        self.search_trends_widget = QWidget()
        self.search_trends_layout = QHBoxLayout()
        self.search_trends_line_edit = QLineEdit()
        self.search_trends_button = QPushButton("Ajouter")
        self.search_trends_layout.addWidget(self.search_trends_line_edit)
        self.search_trends_layout.addWidget(self.search_trends_button)
        self.search_trends_widget.setLayout(self.search_trends_layout)

        # Ajout du tableau ou de la liste d'informations et d'articles récents sur les tendances et les opportunités sélectionnées
        self.trends_table_widget = QTableWidget()
        self.trends_table_widget.setFixedSize(300, 180)
        self.trends_table_widget.setColumnCount(2)
        self.trends_table_widget.setHorizontalHeaderLabels(['Dates', 'Tendances et opportunités sélectionnées'])
        self.trends_table_widget.setRowCount(3)
        self.trends_table_widget.setItem(0, 0, QTableWidgetItem("01/05/2023"))
        self.trends_table_widget.setItem(0, 1, QTableWidgetItem("Tendance 1"))
        self.trends_table_widget.setItem(1, 0, QTableWidgetItem("30/04/2023"))
        self.trends_table_widget.setItem(1, 1, QTableWidgetItem("Tendance 2"))
        self.trends_table_widget.setItem(2, 0, QTableWidgetItem("29/04/2023"))
        self.trends_table_widget.setItem(2, 1, QTableWidgetItem("Tendance 3"))

        # Ajout des filtres pour affiner les résultats en fonction des mots-clés ou des sources d'information
        self.filters_widget = QWidget()
        self.filters_layout = QHBoxLayout()
        self.filters_line_edit = QLineEdit()
        self.filters_button = QPushButton("Filtrer")
        self.filters_layout.addWidget(self.filters_line_edit)
        self.filters_layout.addWidget(self.filters_button)
        self.filters_widget.setLayout(self.filters_layout)

        # Ajout du bouton pour définir des alertes personnalisées en fonction des tendances et des opportunités
        self.alert_button = QPushButton("Définir des alertes personnalisées")

        # Ajout de la section présentant des recommandations d'actions basées sur les tendances et les opportunités identifiées
        self.recommendations_widget = QWidget()
        self.recommendations_layout = QVBoxLayout()
        self.recommendations_label = QLabel("Recommandations d'actions:")
        self.recommendations_text_edit = QTextEdit()
        self.recommendations_layout.addWidget(self.recommendations_label)
        self.recommendations_layout.addWidget(self.recommendations_text_edit)
        self.recommendations_widget.setLayout(self.recommendations_layout)

        # Ajout de tous les éléments à la page de veille stratégique en utilisant un QHBoxLayout pour mieux organiser les widgets
        top_layout = QHBoxLayout()
        top_layout.addWidget(self.trends_list_widget)
        top_layout.addWidget(self.trends_table_widget)
        self.strategic_intel_layout.addLayout(top_layout)

        middle_layout = QHBoxLayout()
        middle_layout.addWidget(self.search_trends_widget)
        middle_layout.addWidget(self.filters_widget)
        self.strategic_intel_layout.addLayout(middle_layout)

        bottom_layout = QVBoxLayout()
        bottom_layout.addWidget(self.alert_button)
        bottom_layout.addWidget(self.recommendations_widget)
        self.strategic_intel_layout.addLayout(bottom_layout)

        self.strategic_intel_widget = QWidget()
        self.strategic_intel_widget.setLayout(self.strategic_intel_layout)
        self.stacked_widget.addWidget(self.strategic_intel_widget)

        # Configuration du layout principal
        self.main_layout = QHBoxLayout()
        self.main_layout.addWidget(self.menu_widget)

        self.vertical_layout = QVBoxLayout()
        self.vertical_layout.addWidget(self.stacked_widget)
        self.main_layout.addLayout(self.vertical_layout)

        # Configuration du widget central
        self.central_widget.setLayout(self.main_layout)
        self.setCentralWidget(self.central_widget)

    def extract_keywords(text, num_keywords=10):
        doc = nlp(text)
        keywords = [token.lemma_ for token in doc if token.is_stop == False and token.is_punct == False]
        keyword_freq = Counter(keywords)
        top_keywords = keyword_freq.most_common(num_keywords)
        return [keyword[0] for keyword in top_keywords]
    
    def extract_country_code(url):
        domain = urlparse(url).netloc.split('.')
        if len(domain) > 1:
            return domain[-1]

        return None

    def extract_article_content(url):
        try:
            response = requests.get(url)
            soup = BeautifulSoup(response.content, "html.parser")
            paragraphs = soup.find_all("p")
            content = "\n".join([p.get_text() for p in paragraphs])
            return content
        except Exception as e:
            print(f"Erreur lors de l'extraction du contenu de l'article : {e}")
            return ""

    search_results_list = [
        # Liste d'articles à analyser...
    ]

    extracted_data = []

    for article in search_results_list:
        publication_date = article.get("date", "")
        title = article.get("title", "")
        link = article.get("link", "")
        snippet = article.get("snippet", "")

        content = extract_article_content(link)
        keywords = extract_keywords(content)

        # Ajoutez le code pour extraire le pays ici
        country_code = extract_country_code(link)
        if country_code:
            country_info = CountryInfo(country_code)
            country_name = country_info.name()
        else:
            country_name = 'Unknown'

        extracted_data.append({
            "date": publication_date,
            "title": title,
            "link": link,
            "snippet": snippet,
            "keywords": keywords,
            "country": country_name  # Ajoutez le pays à la structure de données
        })

    def analyze_articles(self, articles, keywords):
                keyword_data = analyze_keywords(articles, keywords)
                website_data = analyze_websites(articles)
                country_data = analyze_countries(articles)

                return keyword_data, website_data, country_data
    def on_search_button_clicked(self):
        query = self.search_text.text()
        keywords = extract_keywords(query)
        articles = self.perform_search(query)
        keyword_data, website_data, country_data = self.analyze_articles(articles, keywords)
        self.update_visualization_data(keyword_data, website_data, country_data)

    def update_visualization_data(self, keyword_data, website_data, country_data):
            # Mettez à jour les données de visualisation et redessinez les graphiques
            self.matplotlib_widget.update_data(keyword_data, website_data, country_data)

    def perform_search(self, language_codes):
        # Récupérer les options sélectionnées par l'utilisateur
        search_query = self.search_text.text()
        start_date = self.start_date_edit.date().toString("yyyy-MM-dd")
        end_date = self.end_date_edit.date().toString("yyyy-MM-dd")
        country = self.country_combobox.currentText()
        search_query = f"{self.search_text.text()} after:{start_date} before:{end_date}"

        # Utilisez la langue du pays sélectionné comme langue de recherche
        language_code = self.language_codes[country]
        country_code = self.country_codes[country]

        # Récupérer la source de recherche sélectionnée
        selected_source = self.source_combobox.currentText()

        # Récupérer l'état des cases à cocher pour le tri
        sort_relevance = self.sort_relevance_checkbox.isChecked()
        sort_date = self.sort_date_checkbox.isChecked()

         # Définir le paramètre sort en fonction de l'état des cases à cocher
        if sort_relevance:
            sort = "relevance"
        elif sort_date:
            sort = "date"
        else:
            sort = None

        # Définir le paramètre tbm en fonction de la source sélectionnée
        source_options = {
            "Tous": None,
            "Actualités": "nws"
        }

        tbm = source_options[selected_source]

        start_index = 0
        has_results = True

        # Utiliser l'API SerpApi pour effectuer la recherche
        extracted_data = []
        while has_results and start_index < 50:
            api_key = "9b0d4c0366546a7bd81c14d13ae3f304ea744bff2faa67fab9eed518194b7f40"
            params = {
                "q": search_query,
                'start': start_index,
                "tbm": tbm,
                "api_key": api_key,
                "hl": language_code,
                "gl": country_code,
                "google_domain": self.google_domains[country],
                'per_page': 100
            }

            if tbm is not None:
                params["tbm"] = tbm

            response = requests.get("https://serpapi.com/search", params=params)
            search_results = response.json()
            print(search_results)

            # Traitez les résultats de la recherche
            if selected_source == "Actualités":
                results_key = "news_results"
            else:
                results_key = "organic_results"

            search_results_list = search_results.get(results_key, [])

            # Vérifiez s'il y a encore des résultats à récupérer
            if len(search_results_list) == 0:
                break

            start_index += 10  # Incrémentez start_index pour la prochaine page

            # Trier les résultats localement en fonction des options sélectionnées
            if sort_relevance:
                search_results_list = sorted(search_results_list, key=lambda x: x.get("_score", 0), reverse=True)
            elif sort_date:
                search_results_list = sorted(search_results_list, key=lambda x: x.get("date", ""), reverse=True)

            for article in search_results_list:
                publication_date = article.get("date", "")
                title = article.get("title", "")
                link = article.get("link", "")
                snippet = article.get("snippet", "")

                extracted_data.append((publication_date, title, link, snippet))

            # Mettre à jour la barre de progression
            self.progress_bar.setValue(100)

        # Vous pouvez également exporter les résultats au format Excel ou Word, en utilisant les options sélectionnées
        if self.export_excel_checkbox.isChecked():
            self.export_to_excel(extracted_data, search_query="mon_fichier_excel.xlsx")
        if self.export_word_checkbox.isChecked():
            self.export_to_word(extracted_data, search_query="mon_fichier_word.docx")


        # Ajoutez cette ligne à la fin de la méthode perform_search
        return extracted_data

    def create_export_directory(self):
        date_str = datetime.now().strftime("%Y-%m-%d")
        directory = os.path.join("exports", date_str)

        if not os.path.exists(directory):
            os.makedirs(directory)

        return directory

    def export_to_excel(self, data, search_query):
        export_directory = self.create_export_directory()
        timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")

        workbook = Workbook()
        sheet = workbook.active

        # Ajouter les en-têtes de colonnes
        headers = ["Date de publication", "Titre", "Lien", "Résumé"]
        for col_num, header in enumerate(headers, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = header

        # Ajouter les données dans les cellules
        for row_num, row_data in enumerate(data, 2):
            for col_num, cell_data in enumerate(row_data, 1):
                cell = sheet.cell(row=row_num, column=col_num)
                cell.value = cell_data

        # Remplacez les espaces et les caractères spéciaux par des underscores
        sanitized_query = re.sub(r"[^\w\s]", "", search_query).replace(" ", "_")

        # Enregistrer le fichier Excel dans le dossier d'exportation
        workbook.save(os.path.join(export_directory, f"{search_query}_resultats_exportes_{timestamp}.xlsx"))

    def export_to_word(self, data, search_query):
        export_directory = self.create_export_directory()
        timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")
        document = Document()

        # Ajouter les données au document
        for row_data in data:
            date_publication, titre, lien, resume = row_data

            # Ajouter la date de publication
            date_paragraph = document.add_paragraph()
            date_paragraph.add_run(f"Date de publication : {date_publication}").bold = True

            # Ajouter le titre
            document.add_heading(titre, level=1)

            # Ajouter le lien
            document.add_paragraph(f"Lien : {lien}")

            # Ajouter le résumé
            document.add_paragraph(f"Résumé : {resume}")

            # Ajouter un saut de ligne pour séparer les articles
            document.add_paragraph()

        # Remplacez les espaces et les caractères spéciaux par des underscores
        sanitized_query = re.sub(r"[^\w\s]", "", search_query).replace(" ", "_")

        # Enregistrer le fichier Word dans le dossier d'exportation
        document.save(os.path.join(export_directory, f"{search_query}_resultats_exportes_{timestamp}.docx"))

def extract_keywords(text, num_keywords=10):
        doc = nlp(text)
        keywords = [token.lemma_ for token in doc if token.is_stop == False and token.is_punct == False]
        keyword_freq = Counter(keywords)
        top_keywords = keyword_freq.most_common(num_keywords)
        return [keyword[0] for keyword in top_keywords]

def analyze_keywords(articles, keywords):
    keyword_data = {}

    for keyword in keywords:
        keyword_data[keyword] = 0

    for article in articles:
        # Récupérer le lien de l'article
        link = article[2]

        # Récupérer le contenu de la page web
        response = requests.get(link, verify=False)
        soup = BeautifulSoup(response.content, "html.parser")

        # Extraire le texte de l'article
        content = soup.get_text()

        # Analyser les mots-clés dans le contenu
        for keyword in keywords:
            keyword_count = content.lower().count(keyword.lower())
            keyword_data[keyword] += keyword_count

    return keyword_data

def analyze_websites(articles):
    website_count = Counter()

    for article in articles:
        website = article[2]  # Utilisez la clé 'link' au lieu de 'website'
        website_count[website] += 1

    return website_count

def analyze_countries(articles):
    country_count = Counter()

    for article in articles:
        country = article['country']  # Assurez-vous que la clé 'country' existe dans la structure de données
        country_count[country] += 1

    return country_count

if __name__ == "__main__":
    app = QApplication(sys.argv)

    # Configuration de la palette de couleurs
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(53, 53, 53))
    palette.setColor(QPalette.WindowText, Qt.white)
    palette.setColor(QPalette.Base, QColor(25, 25, 25))
    palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
    palette.setColor(QPalette.ToolTipBase, Qt.white)
    palette.setColor(QPalette.ToolTipText, Qt.white)
    palette.setColor(QPalette.Text, Qt.white)
    palette.setColor(QPalette.Button, QColor(53, 53, 53))
    palette.setColor(QPalette.ButtonText, Qt.white)
    palette.setColor(QPalette.BrightText, Qt.red)
    palette.setColor(QPalette.Highlight, QColor(142, 45, 197).lighter())
    palette.setColor(QPalette.HighlightedText, Qt.black)
    app.setPalette(palette)
    app.setStyle("Fusion")
    

    # Affichage de la fenêtre principale
    mainWin = MainWindow()
    mainWin.show()
    sys.exit(app.exec())
