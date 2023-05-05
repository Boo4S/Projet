from __future__ import annotations
import sys
import csv
from typing import Optional
from PyQt5.QtGui import QStandardItemModel, QColor
from PyQt5.QtWidgets import (QApplication, QCheckBox, QComboBox, QDateEdit, QFileDialog, QHBoxLayout, QLabel, QLineEdit, QMainWindow, QMessageBox, QPushButton, QStatusBar, QTableWidget, QTableWidgetItem, QVBoxLayout, QWidget, QGridLayout)
import serpapi
import pandas as pd
from pathlib import Path
from PyQt5.QtGui import QIcon
from docx import Document
from PyQt5.QtCore import QDate
import qtawesome as qta
from PyQt5.QtWidgets import QGraphicsDropShadowEffect


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Ajouter un icône à votre application
        self.setWindowIcon(QIcon(r'C:\Users\Maxim\OneDrive\Bureau\Web scraping\resocom-kyc-logo.png'))

        # Création des widgets pour le menu latéral
        self.menu_widget = QWidget(self)
        self.menu_layout = QVBoxLayout(self.menu_widget)
        self.menu_layout.addWidget(QLabel("Menu"))

        # Création des widgets pour le contenu principal
        self.content_widget = QWidget(self)
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.addWidget(QLabel("Contenu"))

        # Ajout des widgets à la disposition principale de la fenêtre
        main_layout = QHBoxLayout()
        main_layout.addWidget(self.menu_widget)
        main_layout.addWidget(self.content_widget)
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)

        # Configuration de la fenêtre principale
        self.setWindowTitle('RESOSearch')
        self.setGeometry(200, 200, 800, 600)

        # Ajout du thème de couleur personnalisé
        self.setStyleSheet("""
         QMainWindow {
             background-color: #f5f5f5;
         }
         
         QLabel {
             color: #383838;
             font-size: 14px;
         }
         
         QComboBox, QLineEdit, QDateEdit {
             background-color: white;
             border: 1px solid #ccc;
             border-radius: 3px;
             padding: 5px;
         }
         
         QComboBox:focus, QLineEdit:focus, QDateEdit:focus {
             border: 1px solid #6600ff;
         }
         
         QTableWidget {
             background-color: white;
             border: 1px solid #ccc;
             border-radius: 3px;
         }
         
         QHeaderView::section {
             background-color: #6600ff;
             color: white;
             font-size: 12px;
             font-weight: bold;
             padding: 5px;
             border: none;
             border-top-right-radius: 3px;
             border-top-left-radius: 3px;
         }
         
         QPushButton {
             background-color: #6600ff;
             border: none;
             border-radius: 3px;
             color: white;
             font-weight: bold;
             padding: 5px;
             min-width: 100px;
             margin: 10px;
         }
         
         QPushButton:hover {
             background-color: #ff7300;
         }
         
         QStatusBar {
             background-color: #6600ff;
             color: white;
             font-weight: bold;
         }
     """)

        # Ajout des Widgets pour la recherche de résultats
        central_widget = QWidget(self)
        self.setCentralWidget(central_widget)

        vbox = QVBoxLayout()

        # Ajouter une QComboBox pour le choix de thème
        theme_hbox = QHBoxLayout()
        self.theme_label = QLabel('Thème:', central_widget)
        self.theme_combobox = QComboBox(central_widget)
        self.theme_combobox.addItems(['Sécurité informatique', 'Concurrents', 'Autre'])
        theme_hbox.addWidget(self.theme_label)
        theme_hbox.addWidget(self.theme_combobox)
        theme_hbox.addStretch(1)
        vbox.addLayout(theme_hbox)

        self.theme_combobox.currentIndexChanged.connect(self.update_keywords)

        # Dictionnaire de mots-clés pour chaque thème
        self.keywords = {
            'Sécurité informatique': ['cybersécurité', 'hacker', 'pirate informatique'],
            'Concurrents': ['entreprise concurrente', 'comparaison', 'part de marché'],
            'Autre': []  # À définir plus tard
        }
        hbox = QHBoxLayout()
        self.engine_label = QLabel('Moteur de recherche:', central_widget)
        self.engine_combobox = QComboBox(central_widget)
        self.engine_combobox.addItems(['google', 'bing', 'yahoo'])
        hbox.addWidget(self.engine_label)
        hbox.addWidget(self.engine_combobox)
        hbox.addStretch(1)
        vbox.addLayout(hbox)

        hbox = QHBoxLayout()
        self.keywords_label = QLabel('Mots-clés:', central_widget)
        self.keywords_edit = QLineEdit(central_widget)
        hbox.addWidget(self.keywords_label)
        hbox.addWidget(self.keywords_edit)
        hbox.addStretch(1)
        vbox.addLayout(hbox)

        hbox = QHBoxLayout()
        self.start_date_label = QLabel('Date de début:', central_widget)
        self.start_date_edit = QDateEdit(central_widget)
        self.start_date_edit.setCalendarPopup(True)  # Activation du calendrier
        self.start_date_edit.setDisplayFormat('dd/MM/yyyy')  # Format d'affichage
        hbox.addWidget(self.start_date_label)
        hbox.addWidget(self.start_date_edit)
        hbox.addStretch(1)
        vbox.addLayout(hbox)

        hbox = QHBoxLayout()
        self.end_date_label = QLabel('Date de fin:', central_widget)
        self.end_date_edit = QDateEdit(central_widget)
        self.end_date_edit.setCalendarPopup(True)  # Activation du calendrier
        self.end_date_edit.setDisplayFormat('dd/MM/yyyy')  # Format d'affichage
        hbox.addWidget(self.end_date_label)
        hbox.addWidget(self.end_date_edit)
        hbox.addStretch(1)
        vbox.addLayout(hbox)

        hbox = QHBoxLayout()
        self.gl_label = QLabel('Pays ou langue (gl):', central_widget)
        self.gl_edit = QLineEdit(central_widget)
        hbox.addWidget(self.gl_label)
        hbox.addWidget(self.gl_edit)
        hbox.addStretch(1)
        vbox.addLayout(hbox)

        hbox = QHBoxLayout()
        self.source_combobox = QComboBox(central_widget)
        self.source_combobox.addItems(['', 'news', 'blogs', 'web'])
        hbox.addWidget(QLabel('Source :', central_widget))
        hbox.addWidget(self.source_combobox)
        hbox.addStretch(1)
        vbox.addLayout(hbox)

        self.relevant_checkbox = QCheckBox('Trier par pertinence', central_widget)
        vbox.addWidget(self.relevant_checkbox)

        # Ajouter les éléments restants de la recherche de résultats ici :
        self.update_checkbox = QCheckBox('Mettre à jour les résultats existants', central_widget)
        vbox.addWidget(self.update_checkbox)

        hbox = QHBoxLayout()
        self.word_checkbox = QCheckBox('Générer un fichier Word', central_widget)
        self.word_checkbox.setFixedWidth(200)  # Ajout de cette ligne pour fixer la largeur de la case à cocher
        hbox.addWidget(self.word_checkbox)
        hbox.addStretch(1)
        vbox.addLayout(hbox)

        self.search_button = QPushButton('Rechercher', central_widget)
        self.search_button.clicked.connect(self.search_button_clicked)
        vbox.addWidget(self.search_button)
        self.search_button.setIcon(qta.icon('fa5s.search'))
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(10)
        shadow.setOffset(2, 2)
        self.search_button.setGraphicsEffect(shadow)

        # Ajout d'un tableau pour afficher les résultats de recherche
        self.table_widget = QTableWidget(central_widget) 
        self.table_widget.setRowCount(0) 
        self.table_widget.setColumnCount(4) 
        self.table_widget.setHorizontalHeaderLabels(['Titre', 'URL', 'Date', 'Résumé'])

        vbox.addWidget(self.table_widget)

        # Ajout de la barre d'état pour afficher les notifications
        self.statusBar = QStatusBar(self)
        self.setStatusBar(self.statusBar)

        central_widget.setLayout(vbox)

        # Ajout des signaux et des slots
        self.search_button.clicked.connect(self.search_button_clicked)

    def update_keywords(self, index):
        """Met à jour les mots-clés en fonction du thème sélectionné"""
        self.current_theme = self.theme_combobox.currentText()
        self.keywords_edit.setText(' '.join(self.keywords[self.current_theme]))

    def update_status(self, message: str):
        self.statusBar.showMessage(message)

    def create_word_document(self, df, filename):
        doc = Document()

        for index, row in df.iterrows():
            doc.add_heading(row['Titre'], level=1)
            doc.add_heading(row['Date'], level=2)
            doc.add_paragraph(row['Résumé'])
            doc.add_paragraph(row['URL'])

        doc.save(filename)

    def search_button_clicked(self):
    # Récupération des valeurs de recherche
        # Récupération des valeurs de recherche
        engine = self.engine_combobox.currentText()
        keywords = self.keywords_edit.text()
        if not keywords:
            QMessageBox.warning(self, 'Erreur', 'Veuillez saisir des mots-clés')
            
        start_date = self.start_date_edit.date().toString('yyyy-MM-dd')
        end_date = self.end_date_edit.date().toString('yyyy-MM-dd')
        gl = self.gl_edit.text()

        # Création de la liste des filtres sélectionnés
        filters = []
        if self.source_combobox.currentText():
            filters.append(f'source_{self.source_combobox.currentText()}')
        if self.relevant_checkbox.isChecked():
            filters.append('relevant')

        # Extraction des résultats de recherche avec SerpApi
        api_key = '65bd35de40483d3f70e82ed97ba8a86eee6de3fdf4c761ea018f1a133769eb83'
        params = {
            'api_key': api_key,
            'q': f"{keywords} after:{start_date} before:{end_date}",
            'per_page': 100,
            'gl': gl,
            'output': 'json',  # Spécifier le format de sortie JSON pour inclure les mini-résumés
            'sort_by': 'date',  # Trier les résultats par date
            'filters': ','.join(filters),  # Ajouter les filtres sélectionnés
        }

        keyword_list = keywords.split()  # Séparer chaque mot-clé individuel dans une liste
        result_list = []  # Initialiser une liste pour stocker les résultats pour chaque mot-clé individuel

        if engine == 'google':
            search = serpapi.GoogleSearch(params)
        elif engine == 'bing':
            search = serpapi.BingSearch(params)
        elif engine == 'yahoo':
            search = serpapi.YahooSearch(params)
        else:
            QMessageBox.warning(self, 'Erreur', 'Moteur de recherche non valide')

        results = search.get_dict()

        # Conversion des résultats en DataFrame
        df = pd.DataFrame.from_dict(results['organic_results'])

        # Sélection des colonnes à conserver (Titre, URL, Date et Description)
        df = df[['date', 'title', 'link', 'snippet']]

        # Renommer les colonnes pour correspondre à l'affichage de la table widget
        df.columns = ['Date', 'Titre', 'URL', 'Résumé']
        
        # Ajouter les résultats à la liste
        result_list.append(df)

        # Initialisation de la variable "filename"
        filename = f"résultats de recherche {keywords} ({start_date} - {end_date})"

        # Enregistrement des résultats dans un fichier Excel si la case à cocher pour générer un fichier Word n'est pas cochée
        if not self.word_checkbox.isChecked():
            mode = 'w' if not self.update_checkbox.isChecked() else 'a'
            with pd.ExcelWriter(filename + '.xlsx', mode=mode) as writer:
                df.to_excel(writer, sheet_name='Résultats', index=False, startrow=1, startcol=0)
                worksheet = writer.sheets['Résultats']
                worksheet.cell(row=1, column=1).value = df.columns[0]
                worksheet.cell(row=1, column=2).value = df.columns[1]
                worksheet.cell(row=1, column=3).value = df.columns[2]
                worksheet.cell(row=1, column=4).value = df.columns[3]

            self.statusBar.showMessage(f'Le fichier Excel "{filename}.xlsx" a été enregistré avec succès.', 5000)

        # Générer un fichier Word si l'utilisateur a coché la case
        if self.word_checkbox.isChecked():
            self.create_word_document(df, filename + '.docx')
            self.statusBar.showMessage(f'Le fichier Word "{filename}.docx" a été enregistré avec succès.', 5000)

        # Affichage des résultats dans la table widget
        self.table_widget.setRowCount(0)
        for row_number, row_data in df.iterrows():
            title = QTableWidgetItem(row_data['Titre'])
            url = QTableWidgetItem(row_data['URL'])
            date = QTableWidgetItem(str(row_data['Date']))
            summary = QTableWidgetItem(row_data['Résumé'])  # Ajout de la colonne Résumé

            # Insertion des données dans le QTableWidget
            self.table_widget.insertRow(row_number)
            self.table_widget.setItem(row_number, 0, title)
            self.table_widget.setItem(row_number, 1, url)
            self.table_widget.setItem(row_number, 2, date)
            self.table_widget.setItem(row_number, 3, summary)  # Insertion de l'objet QTableWidgetItem

        self.update_status('Les résultats de recherche ont été extraits avec succès')

    def example_usage(self):
        df = pd.DataFrame({'Titre': ['Article 1', 'Article 2'],
                           'Date': ['2022-05-01', '2022-05-02'],
                           'Résumé': ['Résumé de l\'article 1', 'Résumé de l\'article 2'],
                           'URL': ['https://example.com/article1', 'https://example.com/article2']})
        self.create_word_document(df, 'results.docx')

    def open_file(self):
        filename, _ = QFileDialog.getOpenFileName(self, 'Ouvrir le fichier', '', 'CSV(*.csv)')
        if filename:
            df = pd.read_csv(filename)

        # Affichage des résultats dans la table widget
        self.table_widget.setRowCount(0)
        for row_number, row_data in df.iterrows():
            title = QTableWidgetItem(row_data['Titre'])
            url = QTableWidgetItem(row_data['URL'])
            date = QTableWidgetItem(str(row_data['Date']))
            summary = QTableWidgetItem(row_data['Résumé'])

    # Insertion des données dans le QTableWidget
        self.table_widget.insertRow(row_number)
        self.table_widget.setItem(row_number, 0, title)
        self.table_widget.setItem(row_number, 1, url)
        self.table_widget.setItem(row_number, 2, date)
        self.table_widget.setItem(row_number, 3, summary)

        self.update_status(f'Le fichier "{Path(filename).name}" a été ouvert avec succès')

if __name__ == '__main__':
    app = QApplication(sys.argv) 
    window = MainWindow() 
    window.show() 
    sys.exit(app.exec_())

    